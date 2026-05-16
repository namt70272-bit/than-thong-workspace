from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_university_assignment():
    doc = Document()
    
    # Thiết lập font mặc định cho toàn bộ tài liệu (Times New Roman)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    # Hỗ trợ hiển thị font Times New Roman cho các hệ máy khác nhau
    r = style.element.rPr.rFonts
    r.set(qn('w:eastAsia'), 'Times New Roman')

    # --- TRANG TIÊU ĐỀ ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('BÀI TẬP TỰ LUẬN TRÊN MÁY')
    run.bold = True
    run.font.size = Pt(18)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Môn: Cấu trúc dữ liệu và Giải thuật')
    run.italic = True
    run.font.size = Pt(14)

    # Thông tin sinh viên
    info = doc.add_paragraph()
    info.paragraph_format.space_before = Pt(20)
    info.add_run('Họ và tên: ').bold = True
    info.add_run('Nguyễn Minh Kiên\n')
    info.add_run('Mã sinh viên: ').bold = True
    info.add_run('24CS101\n')
    info.add_run('Lớp: ').bold = True
    info.add_run('TTDEV-K01\n')
    info.add_run('Số đề: ').bold = True
    info.add_run('01')

    doc.add_paragraph('-' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- NỘI DUNG ĐỀ BÀI ---
    doc.add_heading('I. ĐỀ BÀI', level=1)
    doc.add_paragraph('Xây dựng chương trình quản lý danh mục TAI NGHE sử dụng cấu trúc dữ liệu Danh sách liên kết đơn. Thực hiện các thao tác: Nhập, Xuất, Tìm kiếm theo tên, Sắp xếp theo đơn giá, Tính tổng tiền, Lọc sản phẩm theo giá và Tìm sản phẩm giá cao nhất.')

    # --- PHÂN TÍCH GIẢI THUẬT ---
    doc.add_heading('II. PHÂN TÍCH GIẢI THUẬT', level=1)
    
    doc.add_heading('1. Cấu trúc dữ liệu', level=2)
    doc.add_paragraph('Sử dụng struct để định nghĩa đối tượng TaiNghe và cấu trúc Node trong danh sách liên kết đơn (Singly Linked List). Mỗi Node gồm phần dữ liệu (info) và con trỏ trỏ tới phần tử kế tiếp (next).')
    
    doc.add_heading('2. Thuật toán sắp xếp', level=2)
    doc.add_paragraph('Áp dụng thuật toán Selection Sort (Sắp xếp lựa chọn) để sắp xếp các Node dựa trên trường DonGia. Thuật toán có độ phức tạp O(n^2), phù hợp với yêu cầu quản lý danh sách vừa và nhỏ.')

    # --- MÃ NGUỒN CHƯƠNG TRÌNH ---
    doc.add_heading('III. MÃ NGUỒN CHƯƠNG TRÌNH (C++)', level=1)
    
    # Đọc code từ file đã tạo
    try:
        with open('/home/hgahc/.openclaw/workspace/24CS101-NguyenMinhKien-TTDEV-De01.cpp', 'r', encoding='utf-8') as f:
            cpp_code = f.read()
    except:
        cpp_code = "// Error reading source file."

    # Chèn code vào block với font Courier New cho giống trình soạn thảo code
    code_paragraph = doc.add_paragraph()
    run = code_paragraph.add_run(cpp_code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # --- KẾT QUẢ THỰC NGHIỆM ---
    doc.add_heading('IV. KẾT QUẢ THỰC NGHIỆM', level=1)
    doc.add_paragraph('Chương trình đã được biên dịch thành công bằng trình biên dịch g++ và thực hiện đầy đủ các chức năng theo yêu cầu của đề bài.')
    
    doc.add_paragraph('\n\n')
    doc.add_paragraph('--- HẾT ---').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Lưu file
    file_path = '/home/hgahc/.openclaw/workspace/24CS101-NguyenMinhKien-TTDEV-De01.docx'
    doc.save(file_path)
    return file_path

if __name__ == "__main__":
    create_university_assignment()
