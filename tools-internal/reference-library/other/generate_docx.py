from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_assignment():
    doc = Document()
    
    # Header
    title = doc.add_heading('BÀI TẬP TỰ LUẬN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Môn: Lý thuyết Xác suất và Thống kê Toán học')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    
    doc.add_paragraph('Người thực hiện: Sinh viên Nguyễn Minh Kiên (Learning Assistant)')
    doc.add_paragraph('Đề số: 01')
    doc.add_paragraph('-' * 20)

    # Bài 1
    doc.add_heading('Bài 1', level=1)
    p1 = doc.add_paragraph()
    p1.add_run('Gọi A1, A2, A3 lần lượt là biến cố phòng chiếu 1, 2, 3 đóng cửa.').italic = True
    doc.add_paragraph('Ta có các xác suất đóng cửa: P(A1) = 0,04; P(A2) = 0,3; P(A3) = 0,15.')
    doc.add_paragraph('Xác suất hoạt động tương ứng: P(not A1) = 0,96; P(not A2) = 0,7; P(not A3) = 0,85.')
    
    doc.add_paragraph('a) Xác suất cả 3 phòng cùng hoạt động:')
    doc.add_paragraph('P(H) = 0,96 * 0,7 * 0,85 = 0,5712.')
    
    doc.add_paragraph('b) Xác suất có ít nhất một phòng chiếu không hoạt động:')
    doc.add_paragraph('Đây là biến cố đối của "Cả 3 phòng cùng hoạt động".')
    doc.add_paragraph('P(K) = 1 - P(H) = 1 - 0,5712 = 0,4288.')

    # Bài 2
    doc.add_heading('Bài 2', level=1)
    doc.add_paragraph('Gọi A là biến cố lấy được quà từ Hộp I: P(A) = 25/30 = 5/6.')
    doc.add_paragraph('Gọi B là biến cố lấy được quà từ Hộp II: P(B) = 15/20 = 3/4.')
    doc.add_paragraph('Sau khi lấy mỗi hộp 1 phiếu, sinh viên chọn 1 phiếu để mở.')
    doc.add_paragraph('Sinh viên nhận được quà khi ít nhất một trong hai phiếu đã lấy là phiếu quà tặng.')
    doc.add_paragraph('Biến cố nhận quà (Q) = A ∪ B.')
    doc.add_paragraph('P(Q) = P(A) + P(B) - P(A ∩ B) = 5/6 + 3/4 - (5/6 * 3/4)')
    doc.add_paragraph('P(Q) = 10/12 + 9/12 - 15/24 = 19/12 - 5/8 = 38/24 - 15/24 = 23/24 ≈ 0,9583.')

    # Bài 3
    doc.add_heading('Bài 3', level=1)
    doc.add_paragraph('Gọi Y là số phiếu tập học viên đã sử dụng. Mỗi lần thi đạt xác suất p = 0,9.')
    doc.add_paragraph('Học viên dừng khi thi đạt 3 lần liên tiếp hoặc hết 5 phiếu.')
    
    doc.add_paragraph('Phân tích các giá trị của Y:')
    doc.add_paragraph('- Y=3: Đạt 3 lần liên tiếp ngay từ đầu. P(Y=3) = 0,9 * 0,9 * 0,9 = 0,729.')
    doc.add_paragraph('- Y=4: Lần 4 đạt và tạo thành chuỗi 3 lần liên tiếp (Trượt-Đạt-Đạt-Đạt). P(Y=4) = 0,1 * 0,9 * 0,9 * 0,9 = 0,0729.')
    doc.add_paragraph('- Y=5: Các trường hợp còn lại. P(Y=5) = 1 - P(Y=3) - P(Y=4) = 1 - 0,729 - 0,0729 = 0,1981.')
    
    doc.add_paragraph('a) Bảng phân phối xác suất của Y:')
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'Y'
    table.cell(0, 1).text = '3'
    table.cell(0, 2).text = '4'
    table.cell(0, 3).text = '5'
    table.cell(1, 0).text = 'P'
    table.cell(1, 1).text = '0,729'
    table.cell(1, 2).text = '0,0729'
    table.cell(1, 3).text = '0,1981'
    
    doc.add_paragraph('b) Thông tin từ bảng phân phối:')
    doc.add_paragraph('- Đa số học viên (72,9%) hoàn thành buổi tập chỉ sau 3 phiếu.')
    doc.add_paragraph('- Số phiếu trung bình kỳ vọng: E(Y) = 3*0,729 + 4*0,0729 + 5*0,1981 = 3,4691 phiếu.')
    
    doc.add_paragraph('c) Hàm phân phối F(y) = P(Y < y):')
    doc.add_paragraph('- Với y ≤ 3: F(y) = 0')
    doc.add_paragraph('- Với 3 < y ≤ 4: F(y) = 0,729')
    doc.add_paragraph('- Với 4 < y ≤ 5: F(y) = 0,729 + 0,0729 = 0,8019')
    doc.add_paragraph('- Với y > 5: F(y) = 1')

    doc.save('/home/hgahc/.openclaw/workspace/BaiTap_XacSuat_De1.docx')

if __name__ == "__main__":
    create_assignment()
